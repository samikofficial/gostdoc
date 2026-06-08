"""Главный инвариант: текст тела до и после обработки идентичен (ТЗ, разд. 3, 10, 8.15)."""

from __future__ import annotations

import pytest
from docx import Document

from gostdoc.formatter import extract_body_text, format_document

# Все валидные .docx-фикстуры (без broken/legacy и без already_gost — он сам выход форматтера).
VALID_FIXTURES = [
    "calibri_default.docx",
    "direct_formatting.docx",
    "multi_section.docx",
    "styled_headings.docx",
    "with_table.docx",
    "with_image.docx",
    "with_list.docx",
    "localized_styles.docx",
    "title_and_body.docx",
]


@pytest.mark.parametrize("name", VALID_FIXTURES)
def test_body_text_unchanged(name, fixtures_dir, tmp_path):
    src = fixtures_dir / name
    before = extract_body_text(Document(str(src)))
    out = tmp_path / f"out_{name}"
    format_document(str(src), str(out))
    after = extract_body_text(Document(str(out)))
    assert before == after


def test_invariant_proves_non_empty(fixtures_dir):
    # Защита от ложно-зелёного: текст тела действительно непуст.
    text = extract_body_text(Document(str(fixtures_dir / "calibri_default.docx")))
    assert text.strip()
