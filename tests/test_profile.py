"""Тесты профилей: переопределения под методички (поля, № страницы, полужирность)."""

from __future__ import annotations

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from gostdoc.classify import CATEGORY_HEADING, classify_paragraph
from gostdoc.cli import EXIT_ERROR, EXIT_NONCOMPLIANT, EXIT_OK, main
from gostdoc.formatter import check_compliance, format_document
from gostdoc.profile import GOST, build_profile, parse_margins
from gostdoc.sections import normalize_margins, setup_page_numbering
from gostdoc.xml_utils import lengths_equal, paragraph_has_page_field


def test_parse_margins_ok():
    left, right, top, bottom = parse_margins("30,10,20,20")
    assert lengths_equal(left, Mm(30))
    assert lengths_equal(right, Mm(10))
    assert lengths_equal(top, Mm(20))
    assert lengths_equal(bottom, Mm(20))


@pytest.mark.parametrize("bad", ["30,10,20", "30,10,20,20,5", "a,b,c,d", ""])
def test_parse_margins_invalid(bad):
    with pytest.raises(ValueError):
        parse_margins(bad)


def test_build_profile_default_is_gost():
    assert build_profile() == GOST


def test_custom_margins_applied(load_fixture):
    doc = load_fixture("multi_section.docx")
    normalize_margins(doc, build_profile(margins="25,25,25,25"))
    for s in doc.sections:
        assert lengths_equal(s.left_margin, Mm(25))
        assert lengths_equal(s.right_margin, Mm(25))


def test_page_number_bottom_right(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc, build_profile(page_number="bottom-right"))
    footer = doc.sections[0].footer
    page_paras = [p for p in footer.paragraphs if paragraph_has_page_field(p)]
    assert page_paras
    assert all(p.alignment == WD_ALIGN_PARAGRAPH.RIGHT for p in page_paras)


def test_page_number_top_right(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc, build_profile(page_number="top-right"))
    first = doc.sections[0]
    assert any(paragraph_has_page_field(p) for p in first.header.paragraphs)
    assert not any(paragraph_has_page_field(p) for p in first.footer.paragraphs)


def test_page_number_none(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc, build_profile(page_number="none"))
    first = doc.sections[0]
    assert not any(paragraph_has_page_field(p) for p in first.footer.paragraphs)
    assert not any(paragraph_has_page_field(p) for p in first.header.paragraphs)


def test_no_bold_headings(fixtures_dir, tmp_path):
    out = tmp_path / "o.docx"
    format_document(
        str(fixtures_dir / "styled_headings.docx"), str(out), profile=build_profile(bold_headings=False)
    )
    doc = Document(str(out))
    headings = [p for p in doc.paragraphs if classify_paragraph(p, in_table=False) == CATEGORY_HEADING]
    assert headings
    for p in headings:
        for r in p.runs:
            assert r.bold is not True


def test_check_uses_profile_margins(fixtures_dir, tmp_path):
    profile = build_profile(margins="20,20,20,20")
    out = tmp_path / "o.docx"
    format_document(str(fixtures_dir / "calibri_default.docx"), str(out), profile=profile)
    doc = Document(str(out))
    assert check_compliance(doc, profile) == []  # соответствует своему профилю
    assert check_compliance(doc, GOST)  # но не базовому ГОСТ (поля иные)


def test_cli_margins_and_page_number(fixtures_dir, tmp_path):
    out = tmp_path / "o.docx"
    rc = main(
        [
            str(fixtures_dir / "title_and_body.docx"),
            "-o",
            str(out),
            "--margins",
            "20,20,20,20",
            "--page-number",
            "bottom-right",
            "--page-number-size",
            "12",
        ]
    )
    assert rc == EXIT_OK
    doc = Document(str(out))
    assert lengths_equal(doc.sections[0].right_margin, Mm(20))
    footer = doc.sections[0].footer
    page_paras = [p for p in footer.paragraphs if paragraph_has_page_field(p)]
    assert page_paras and page_paras[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # кегль номера страницы — 12
    assert any(r.font.size == Pt(12) for p in footer.paragraphs for r in p.runs)


def test_cli_bad_margins_returns_error(fixtures_dir, capsys):
    rc = main([str(fixtures_dir / "calibri_default.docx"), "--margins", "oops"])
    assert rc == EXIT_ERROR
    assert "поля" in capsys.readouterr().err.lower()


def test_cli_check_with_profile(fixtures_dir, tmp_path):
    out = tmp_path / "o.docx"
    format_document(
        str(fixtures_dir / "calibri_default.docx"), str(out), profile=build_profile(margins="20,20,20,20")
    )
    # проверка с тем же профилем — соответствует
    assert main([str(out), "--check", "--margins", "20,20,20,20"]) == EXIT_OK
    # проверка по базовому ГОСТ — несоответствие
    assert main([str(out), "--check"]) == EXIT_NONCOMPLIANT
