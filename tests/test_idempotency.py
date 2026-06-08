"""Идемпотентность: повторный прогон не меняет проверяемых свойств (ТЗ, разд. 10, 8.14)."""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from gostdoc.formatter import extract_body_text, format_document
from gostdoc.xml_utils import lengths_equal

FIXTURES = ["calibri_default.docx", "multi_section.docx", "title_and_body.docx", "with_table.docx"]


def _count_page_fields(document) -> int:
    """Поля PAGE только в собственных (не слинкованных) колонтитулах — иначе общий
    колонтитул слинкованных секций считается несколько раз."""
    n = 0
    for section in document.sections:
        footer = section.footer
        if footer.is_linked_to_previous:
            continue
        for p in footer.paragraphs:
            for instr in p._p.findall(".//" + qn("w:instrText")):
                if "PAGE" in (instr.text or "").upper():
                    n += 1
    return n


@pytest.mark.parametrize("name", FIXTURES)
def test_double_run_is_stable(name, fixtures_dir, tmp_path):
    src = fixtures_dir / name
    out1 = tmp_path / f"once_{name}"
    out2 = tmp_path / f"twice_{name}"
    format_document(str(src), str(out1))
    format_document(str(out1), str(out2))

    d1, d2 = Document(str(out1)), Document(str(out2))

    # Текст тела стабилен.
    assert extract_body_text(d1) == extract_body_text(d2)
    # Поле PAGE не задвоилось: ровно одно в собственном колонтитуле первой секции.
    assert _count_page_fields(d1) == _count_page_fields(d2) == 1
    # Поля секций совпадают.
    for s1, s2 in zip(d1.sections, d2.sections):
        assert lengths_equal(s1.left_margin, s2.left_margin)
        assert lengths_equal(s1.right_margin, s2.right_margin)
        assert lengths_equal(s1.top_margin, s2.top_margin)
        assert lengths_equal(s1.bottom_margin, s2.bottom_margin)
