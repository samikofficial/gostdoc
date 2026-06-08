"""Тесты категоризации абзацев (ТЗ, разд. 7.2)."""

from __future__ import annotations

from docx import Document

from gostdoc.classify import (
    CATEGORY_BODY,
    CATEGORY_EMPTY,
    CATEGORY_HEADING,
    CATEGORY_LIST_ITEM,
    CATEGORY_TABLE_CELL,
    classify_paragraph,
)


def test_plain_body_paragraphs(load_fixture):
    doc = load_fixture("calibri_default.docx")
    for p in doc.paragraphs:
        assert classify_paragraph(p, in_table=False) == CATEGORY_BODY


def test_styled_headings_are_headings(load_fixture):
    doc = load_fixture("styled_headings.docx")
    cats = [classify_paragraph(p, in_table=False) for p in doc.paragraphs]
    assert cats[0] == CATEGORY_HEADING  # Heading 1
    assert cats[1] == CATEGORY_BODY
    assert cats[2] == CATEGORY_HEADING  # Heading 2
    assert cats[3] == CATEGORY_BODY


def test_localized_heading_and_body(load_fixture):
    doc = load_fixture("localized_styles.docx")
    cats = [classify_paragraph(p, in_table=False) for p in doc.paragraphs]
    assert cats[0] == CATEGORY_HEADING  # «Заголовок 1»
    assert cats[1] == CATEGORY_BODY  # «Обычный»


def test_list_items_detected_via_style_numpr(load_fixture):
    doc = load_fixture("with_list.docx")
    cats = [classify_paragraph(p, in_table=False) for p in doc.paragraphs]
    assert cats[0] == CATEGORY_BODY  # «Перечень пунктов:»
    assert cats[1:] == [CATEGORY_LIST_ITEM] * (len(cats) - 1)


def test_table_cell_paragraphs(load_fixture):
    doc = load_fixture("with_table.docx")
    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                assert classify_paragraph(p, in_table=True) == CATEGORY_TABLE_CELL


def test_empty_paragraph():
    doc = Document()
    p = doc.add_paragraph("")
    assert classify_paragraph(p, in_table=False) == CATEGORY_EMPTY


def test_image_paragraph_with_text_is_body(load_fixture):
    doc = load_fixture("with_image.docx")
    # Абзац содержит текст и рисунок — это не пустой абзац.
    assert classify_paragraph(doc.paragraphs[0], in_table=False) == CATEGORY_BODY
