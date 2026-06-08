"""Тесты устойчивости на вырожденных документах (ТЗ, подв. камень №21)."""

from __future__ import annotations

import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm

import make_fixtures
from gostdoc.formatter import format_document


def test_empty_document_does_not_crash(tmp_path):
    src = tmp_path / "empty.docx"
    Document().save(str(src))
    out = tmp_path / "empty.gost.docx"
    format_document(str(src), str(out))
    assert out.exists()


def test_omml_formula_preserved(fixtures_dir, tmp_path):
    src = fixtures_dir / "with_formula.docx"
    omath = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"
    before = len(list(Document(str(src)).element.iter(omath)))
    assert before >= 1, "в фикстуре нет формулы OMML"
    out = tmp_path / "formula.gost.docx"
    format_document(str(src), str(out), detect_structure=True)
    reopened = Document(str(out))
    after = len(list(reopened.element.iter(omath)))
    assert after == before  # формула сохранена, не сломана нормализацией


def test_image_only_paragraph_keeps_image(tmp_path):
    doc = Document()
    doc.add_paragraph().add_run().add_picture(io.BytesIO(make_fixtures._PNG_1x1), width=Mm(20))
    src = tmp_path / "img.docx"
    doc.save(str(src))
    out = tmp_path / "img.gost.docx"
    format_document(str(src), str(out))
    reopened = Document(str(out))
    assert reopened.paragraphs[0]._p.findall(".//" + qn("w:drawing"))
