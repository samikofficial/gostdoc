"""Тесты нумерации страниц: поле PAGE по центру, титул без номера, сквозная (ТЗ, DoD + подв. камни №4,5,14)."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from gostdoc.sections import setup_page_numbering
from gostdoc.xml_utils import paragraph_has_page_field


def _count_page_fields(footer) -> int:
    n = 0
    for p in footer.paragraphs:
        for instr in p._p.findall(".//" + qn("w:instrText")):
            if "PAGE" in (instr.text or "").upper():
                n += 1
    return n


def test_primary_footer_has_centered_page_field(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc)
    first = doc.sections[0]
    assert first.different_first_page_header_footer is True
    footer = first.footer
    page_paras = [p for p in footer.paragraphs if paragraph_has_page_field(p)]
    assert page_paras, "в основном колонтитуле нет поля PAGE"
    assert all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in page_paras)


def test_first_page_has_no_number(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc)
    fpf = doc.sections[0].first_page_footer
    assert not any(paragraph_has_page_field(p) for p in fpf.paragraphs)


def test_page_numbering_is_idempotent(load_fixture):
    doc = load_fixture("title_and_body.docx")
    setup_page_numbering(doc)
    setup_page_numbering(doc)
    assert _count_page_fields(doc.sections[0].footer) == 1


def test_subsequent_sections_continue_numbering(load_fixture):
    doc = load_fixture("multi_section.docx")
    setup_page_numbering(doc)
    assert doc.sections[1].footer.is_linked_to_previous is True
