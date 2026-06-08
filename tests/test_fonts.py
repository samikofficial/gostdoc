"""Тесты нормализации run'ов: гарнитура/кегль/цвет + сохранение акцентов (ТЗ, разд. 8.1, 8.2, 8.6)."""

from __future__ import annotations

from docx.oxml.ns import qn
from docx.shared import Pt

from docx import Document

from gostdoc import constants as c
from gostdoc.formatter import format_document
from gostdoc.runs import normalize_run


def _rfonts_slot(run, slot: str):
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return None
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        return None
    return rf.get(qn("w:" + slot))


def _color_val(run):
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return None
    color = rpr.find(qn("w:color"))
    if color is None:
        return None
    return color.get(qn("w:val"))


def _accents(run):
    return (run.bold, run.italic, run.underline, run.font.superscript, run.font.subscript)


def test_normalize_sets_font_size_color_all_slots(load_fixture):
    doc = load_fixture("direct_formatting.docx")
    for p in doc.paragraphs:
        for run in p.runs:
            normalize_run(run)
    for p in doc.paragraphs:
        for run in p.runs:
            assert run.font.name == c.FONT_NAME
            assert _rfonts_slot(run, "ascii") == c.FONT_NAME
            assert _rfonts_slot(run, "hAnsi") == c.FONT_NAME
            assert _rfonts_slot(run, "cs") == c.FONT_NAME
            assert run.font.size == Pt(14)
            assert _color_val(run) == "auto"


def test_normalize_preserves_accents(load_fixture):
    doc = load_fixture("direct_formatting.docx")
    runs = [r for p in doc.paragraphs for r in p.runs]
    before = [_accents(r) for r in runs]
    for r in runs:
        normalize_run(r)
    after = [_accents(r) for r in runs]
    assert before == after
    # Убедимся, что в фикстуре реально были акценты (иначе тест ничего не проверяет).
    flat = [flag for tup in before for flag in tup]
    assert any(flag is True for flag in flat)


def test_normalize_does_not_recreate_run_keeps_drawing(load_fixture):
    doc = load_fixture("with_image.docx")
    p = doc.paragraphs[0]
    drawings_before = len(p._p.findall(".//" + qn("w:drawing")))
    for run in p.runs:
        normalize_run(run)
    drawings_after = len(p._p.findall(".//" + qn("w:drawing")))
    assert drawings_before == drawings_after == 1


def test_hyperlink_run_font_normalized(fixtures_dir, tmp_path):
    out = tmp_path / "hl.docx"
    format_document(str(fixtures_dir / "with_hyperlink.docx"), str(out))
    p = Document(str(out)).paragraphs[0]
    fonts = [r.font.name for hl in p.hyperlinks for r in hl.runs]
    assert fonts, "в фикстуре нет run'ов внутри гиперссылки"
    assert all(name == c.FONT_NAME for name in fonts)
