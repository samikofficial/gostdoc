"""Тесты интервалов/отступов/выравнивания по категориям (ТЗ, разд. 7.2, 8.10, 8.11)."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from gostdoc import constants as c
from gostdoc.classify import (
    CATEGORY_BODY,
    CATEGORY_CAPTION,
    CATEGORY_LIST_ITEM,
    CATEGORY_TABLE_CELL,
)
from gostdoc.paragraphs import apply_paragraph_format
from gostdoc.xml_utils import lengths_equal


def test_body_paragraph_format(load_fixture):
    doc = load_fixture("calibri_default.docx")
    for p in doc.paragraphs:
        apply_paragraph_format(p, CATEGORY_BODY)
    for p in doc.paragraphs:
        pf = p.paragraph_format
        assert pf.line_spacing == c.LINE_SPACING_BODY
        # python-docx читает line=360/auto как ONE_POINT_FIVE — это полуторный по ГОСТ.
        assert pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
        assert pf.space_before == Pt(0)
        assert pf.space_after == Pt(0)
        assert lengths_equal(pf.first_line_indent, Cm(1.25))
        assert pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_table_cell_format_no_indent_single_spacing(load_fixture):
    doc = load_fixture("with_table.docx")
    cell_paras = [p for row in doc.tables[0].rows for cell in row.cells for p in cell.paragraphs]
    for p in cell_paras:
        apply_paragraph_format(p, CATEGORY_TABLE_CELL)
    for p in cell_paras:
        pf = p.paragraph_format
        assert lengths_equal(pf.first_line_indent, Cm(0))
        assert pf.line_spacing_rule == WD_LINE_SPACING.SINGLE
        # Выравнивание ячеек не навязываем — остаётся неустановленным.
        assert pf.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY


def test_list_item_keeps_first_line_indent(load_fixture):
    doc = load_fixture("with_list.docx")
    list_paras = doc.paragraphs[1:]  # первый — вводный body-абзац
    for p in list_paras:
        before_indent = p.paragraph_format.first_line_indent
        apply_paragraph_format(p, CATEGORY_LIST_ITEM)
        # Отступ первой строки списка не трогаем (подводный камень №10).
        assert p.paragraph_format.first_line_indent == before_indent
        assert p.paragraph_format.line_spacing == c.LINE_SPACING_BODY


def test_caption_no_indent_single_spacing():
    from docx import Document

    doc = Document()
    p = doc.add_paragraph("Рисунок 1 — Схема")
    p.paragraph_format.first_line_indent = Cm(1.25)
    apply_paragraph_format(p, CATEGORY_CAPTION)
    assert lengths_equal(p.paragraph_format.first_line_indent, Cm(0))
    assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
