"""Тесты Фазы 2: распознавание неразмеченной структуры (ТЗ, разд. 12).

Фикстура unmarked_structure кодирует реальные паттерны и ложные срабатывания,
найденные на 5 настоящих ВКР.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from gostdoc.classify import CATEGORY_HEADING, classify_paragraph
from gostdoc.cli import EXIT_OK, main
from gostdoc.detect import detect_and_mark
from gostdoc.formatter import extract_body_text, format_document


def _by_text(doc, needle):
    for p in doc.paragraphs:
        if p.text.strip().startswith(needle):
            return p
    raise AssertionError(f"не найден абзац: {needle!r}")


def _exact(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise AssertionError(f"не найден абзац с точным текстом: {text!r}")


def test_structural_elements_detected(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    for word in ("ВВЕДЕНИЕ", "Заключение", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = _exact(doc, word)
        assert classify_paragraph(p, in_table=False) == CATEGORY_HEADING
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert p.paragraph_format.page_break_before is True
        # Прописные через w:caps, текст не изменён.
        assert any(r._r.find(qn("w:rPr")).find(qn("w:caps")) is not None for r in p.runs)


def test_numbered_multilevel_headings_detected(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    assert _by_text(doc, "1.1 История возникновения").style.name == "Heading 2"
    assert _by_text(doc, "1.1.1 Подробности").style.name == "Heading 3"
    # Реальный провал doc06: номер без пробела после точки, уровень считаем верно.
    assert _by_text(doc, "3.2.Особенности").style.name == "Heading 2"
    # Реальный провал doc05: глава с римским номером → Heading 1.
    assert _by_text(doc, "ГЛАВА II. Экспериментальная").style.name == "Heading 1"
    # Реальный провал doc08: одноуровневая полужирная глава → Heading 1.
    assert _by_text(doc, "1 Теоретические основы исследуемого").style.name == "Heading 1"
    # Реальный провал doc08: длинный (>100) многоуровневый подзаголовок → Heading 2.
    assert _by_text(doc, "2.1 Система показателей").style.name == "Heading 2"


def test_single_level_requires_bold(load_fixture):
    # Одноуровневые НЕ полужирные «1. …» (поле/пункт списка) заголовками не становятся.
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    assert _by_text(doc, "1. Тема:").style.name == "Normal"
    assert _by_text(doc, "1. Кубики Кооса").style.name == "Normal"


def test_structural_promoted_over_existing_heading(load_fixture):
    # «СОДЕРЖАНИЕ» в источнике помечено Heading 3 — должно стать структурным (центр, Heading 1).
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    p = _exact(doc, "СОДЕРЖАНИЕ")
    assert p.style.name == "Heading 1"
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_numbered_structural_element_detected(load_fixture):
    # Реальный провал doc02: «5. Список использованных источников» с ведущим номером.
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    p = _by_text(doc, "5. Список использованных источников")
    assert classify_paragraph(p, in_table=False) == CATEGORY_HEADING
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_false_positives_not_marked(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    # Поле-задание и пункт списка (одноуровневые «1.») — НЕ заголовки.
    assert _by_text(doc, "1. Тема:").style.name == "Normal"
    assert _by_text(doc, "1. Кубики Кооса").style.name == "Normal"
    # Обычный абзац тела — не тронут.
    assert _by_text(doc, "Текст введения").style.name == "Normal"


def test_toc_lines_protected_not_marked(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    # Строки оглавления (лидеры + номер страницы) не должны стать заголовками.
    assert _by_text(doc, "ВВЕДЕНИЕ ......").style.name == "Normal"
    assert _by_text(doc, "1.1 История вопроса .....").style.name == "Normal"
    # Реальный провал из doc06: лидеры разбиты пробелами — тоже должно защищаться.
    assert _by_text(doc, "Раздел 2. Проблема интертекста").style.name == "Normal"
    # Реальный провал из doc04: короткий лидер из двух точек («..20»).
    assert _by_text(doc, "1.2.1. Основные подходы к изучению").style.name == "Normal"
    # Реальный провал из web01/web02: TOC-строка с табом-лидером + номер.
    assert _by_text(doc, "Глава 5. Название с табуляцией").style.name == "Normal"


def test_captions_detected(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    fig = _by_text(doc, "Рисунок 1 — Схема")
    fig_abbr = _by_text(doc, "Рис. 1.1.")
    tbl = _by_text(doc, "Таблица 2 —")
    assert fig.style.name == "Caption"
    assert fig.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert fig_abbr.style.name == "Caption"
    assert tbl.style.name == "Caption"
    assert tbl.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_figure_reference_in_text_not_caption(load_fixture):
    doc = load_fixture("unmarked_structure.docx")
    detect_and_mark(doc)
    # «Рисунок 1 показывает …» — ссылка в тексте, не подпись.
    assert _by_text(doc, "Рисунок 1 показывает").style.name == "Normal"


def test_detect_preserves_text_invariant(fixtures_dir, tmp_path):
    src = fixtures_dir / "unmarked_structure.docx"
    before = extract_body_text(Document(str(src)))
    out = tmp_path / "out.docx"
    format_document(str(src), str(out), detect_structure=True)
    after = extract_body_text(Document(str(out)))
    assert before == after  # caps — это отображение, текст не меняется


def test_detect_log_returned(fixtures_dir, tmp_path):
    out = tmp_path / "out.docx"
    messages = format_document(str(fixtures_dir / "unmarked_structure.docx"), str(out), detect_structure=True)
    assert any("Фаза 2" in m for m in messages)


def test_library_default_does_not_detect(fixtures_dir, tmp_path):
    # Библиотечный format_document по умолчанию консервативен (детекция выключена).
    out = tmp_path / "out.docx"
    format_document(str(fixtures_dir / "unmarked_structure.docx"), str(out))
    doc = Document(str(out))
    assert _by_text(doc, "1.1 История возникновения").style.name == "Normal"


def test_cli_detects_by_default(fixtures_dir, tmp_path):
    # CLI по умолчанию включает детекцию.
    out = tmp_path / "o.docx"
    assert main([str(fixtures_dir / "unmarked_structure.docx"), "-o", str(out)]) == EXIT_OK
    doc = Document(str(out))
    assert _by_text(doc, "1.1 История возникновения").style.name == "Heading 2"


def test_cli_no_detect_flag_disables(fixtures_dir, tmp_path):
    out = tmp_path / "o.docx"
    rc = main([str(fixtures_dir / "unmarked_structure.docx"), "-o", str(out), "--no-detect-structure"])
    assert rc == EXIT_OK
    doc = Document(str(out))
    assert _by_text(doc, "1.1 История возникновения").style.name == "Normal"
