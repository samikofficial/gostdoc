"""Генератор «грязных» .docx-фикстур для тестов (ТЗ, разд. 11).

Каждая фикстура моделирует типовой дефект студенческого оформления.
Запуск: python tests/make_fixtures.py  (или вызывается из conftest при отсутствии файлов).

Намеренно НЕ зависит от gostdoc, кроме already_gost.docx, который является
выходом форматтера (генерируется лениво, если форматтер уже реализован) — так
already_gost всегда согласован с текущей логикой и мы не дублируем XML-код вставки
поля PAGE.
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

FIXTURES_DIR = Path(__file__).parent / "fixtures"


def _make_png_1x1() -> bytes:
    """Валидный PNG 1x1 (RGB) из stdlib — для встроенного рисунка фикстуры."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit, RGB
    raw = b"\x00\xff\x00\x00"  # filter byte 0 + один пиксель RGB
    idat = zlib.compress(raw)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


_PNG_1x1 = _make_png_1x1()


def _set_run(run, name: str, size_pt: float) -> None:
    """Прямое форматирование run'а — имитация ручной правки студентом."""
    run.font.name = name
    run.font.size = Pt(size_pt)


def calibri_default(path: Path) -> None:
    """Calibri 11, одинарный интервал, без отступа, по левому краю."""
    doc = Document()
    for text in (
        "Первый абзац основного текста, оформленный как попало.",
        "Второй абзац с тем же дефолтным оформлением Calibri.",
        "Третий абзац для объёма.",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Pt(0)
        _set_run(p.add_run(text), "Calibri", 11)
    doc.save(str(path))


def direct_formatting(path: Path) -> None:
    """Прямое форматирование run'ов: разные шрифты/кегли/цвета + b/i/u/sup/sub."""
    doc = Document()
    p = doc.add_paragraph()
    _set_run(p.add_run("Обычный Arial 12. "), "Arial", 12)
    r_bold = p.add_run("Полужирный. ")
    _set_run(r_bold, "Calibri", 11)
    r_bold.bold = True
    r_it = p.add_run("Курсив. ")
    _set_run(r_it, "Verdana", 13)
    r_it.italic = True
    r_un = p.add_run("Подчёркнутый. ")
    _set_run(r_un, "Calibri", 11)
    r_un.underline = True
    r_color = p.add_run("Цветной. ")
    _set_run(r_color, "Arial", 12)
    r_color.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p2 = doc.add_paragraph()
    _set_run(p2.add_run("Формула H"), "Calibri", 11)
    r_sub = p2.add_run("2")
    _set_run(r_sub, "Calibri", 11)
    r_sub.font.subscript = True
    _set_run(p2.add_run("O и степень x"), "Calibri", 11)
    r_sup = p2.add_run("2")
    _set_run(r_sup, "Calibri", 11)
    r_sup.font.superscript = True
    _set_run(p2.add_run(" в тексте."), "Calibri", 11)
    doc.save(str(path))


def multi_section(path: Path) -> None:
    """Две секции с разными полями (проверка обхода ВСЕХ секций)."""
    doc = Document()
    s0 = doc.sections[0]
    s0.left_margin = Mm(25)
    s0.right_margin = Mm(25)
    s0.top_margin = Mm(25)
    s0.bottom_margin = Mm(25)
    doc.add_paragraph("Текст первой секции.")

    new = doc.add_section(WD_SECTION.NEW_PAGE)
    new.left_margin = Mm(10)
    new.right_margin = Mm(10)
    new.top_margin = Mm(10)
    new.bottom_margin = Mm(10)
    doc.add_paragraph("Текст второй секции с другими полями.")
    doc.save(str(path))


def styled_headings(path: Path) -> None:
    """Заголовки размечены стилями Heading 1/2 (но без полужирного — его наводит форматтер)."""
    doc = Document()
    h1 = doc.add_paragraph("Введение", style="Heading 1")
    for r in h1.runs:
        r.bold = False
        r.underline = True  # дефект: подчёркивание, которое форматтер должен снять
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("Обзор литературы", style="Heading 2")
    doc.add_paragraph("Текст обзора.")
    doc.save(str(path))


def with_table(path: Path) -> None:
    """Таблица с текстом в ячейках + абзацы тела."""
    doc = Document()
    doc.add_paragraph("Абзац перед таблицей.")
    table = doc.add_table(rows=2, cols=2)
    cells = [("Заголовок A", "Заголовок B"), ("Значение 1", "Значение 2")]
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.paragraphs[0].add_run(cells[r][c])
    doc.add_paragraph("Абзац после таблицы.")
    doc.save(str(path))


def with_image(path: Path) -> None:
    """Встроенный (inline) рисунок внутри run — проверка, что run не пересоздаётся."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Текст до рисунка. ")
    run = p.add_run()
    run.add_picture(io.BytesIO(_PNG_1x1), width=Mm(20))
    p.add_run(" Текст после рисунка.")
    doc.save(str(path))


def with_list(path: Path) -> None:
    """Нумерованный и маркированный списки (numPr на уровне стиля)."""
    doc = Document()
    doc.add_paragraph("Перечень пунктов:")
    for t in ("Первый пункт", "Второй пункт", "Третий пункт"):
        doc.add_paragraph(t, style="List Number")
    for t in ("Маркер один", "Маркер два"):
        doc.add_paragraph(t, style="List Bullet")
    doc.save(str(path))


def localized_styles(path: Path) -> None:
    """Русские имена стилей (Обычный, Заголовок 1) — проверка карты синонимов."""
    doc = Document()
    styles = doc.styles
    normal_ru = styles.add_style("Обычный текст", WD_STYLE_TYPE.PARAGRAPH)
    normal_ru.base_style = styles["Normal"]
    heading_ru = styles.add_style("Заголовок 1 (рус)", WD_STYLE_TYPE.PARAGRAPH)
    heading_ru.base_style = styles["Normal"]

    # Переименовываем сами стили на русские имена, по которым их и будут искать.
    from docx.oxml.ns import qn

    normal_ru._element.find(qn("w:name")).set(qn("w:val"), "Обычный")
    heading_ru._element.find(qn("w:name")).set(qn("w:val"), "Заголовок 1")

    h = doc.add_paragraph("Раздел документа", style="Заголовок 1")
    _set_run_existing(h, "Calibri", 11)
    b = doc.add_paragraph("Текст раздела.", style="Обычный")
    _set_run_existing(b, "Calibri", 11)
    doc.save(str(path))


def _set_run_existing(paragraph, name: str, size_pt: float) -> None:
    for r in paragraph.runs:
        r.font.name = name
        r.font.size = Pt(size_pt)


def title_and_body(path: Path) -> None:
    """Титульная страница + несколько страниц тела (для нумерации/титула)."""
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(title.add_run("МИНИСТЕРСТВО ОБРАЗОВАНИЯ"), "Calibri", 14)
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(author.add_run("Курсовая работа\nИванов И. И."), "Calibri", 12)
    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(year.add_run("Москва 2026"), "Calibri", 12)

    doc.add_page_break()
    for i in range(1, 4):
        _set_run(doc.add_paragraph().add_run(f"Абзац тела номер {i} на странице тела."), "Calibri", 11)
    doc.save(str(path))


def already_gost(path: Path) -> bool:
    """Документ, уже соответствующий ГОСТ. Это выход форматтера на простом входе.

    Генерируется только если форматтер уже реализован (иначе пропускается —
    идемпотентность проверяется и без него через process(process(x))).
    Возвращает True, если файл создан.
    """
    try:
        from gostdoc.formatter import format_document
    except Exception:
        return False
    src = FIXTURES_DIR / "calibri_default.docx"
    if not src.exists():
        calibri_default(src)
    format_document(str(src), str(path))
    return True


def with_hyperlink(path: Path) -> None:
    """Гиперссылка с прямым форматированием Calibri внутри (подв. камень №17)."""
    doc = Document()
    p = doc.add_paragraph()
    _set_run(p.add_run("См. ресурс: "), "Calibri", 11)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "ref1")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11 пт
    rpr.append(rfonts)
    rpr.append(sz)
    t = OxmlElement("w:t")
    t.text = "перейти по ссылке"
    r.append(rpr)
    r.append(t)
    hyperlink.append(r)
    p._p.append(hyperlink)
    _set_run(p.add_run(" — конец абзаца."), "Calibri", 11)
    doc.save(str(path))


def with_revisions(path: Path) -> None:
    """Документ с исправлениями (tracked changes) — для предупреждения (подв. камень №19)."""
    doc = Document()
    p = doc.add_paragraph()
    _set_run(p.add_run("Текст с правкой рецензента. "), "Calibri", 11)
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Рецензент")
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "вставленный фрагмент"
    r.append(t)
    ins.append(r)
    p._p.append(ins)
    doc.save(str(path))


def mirror_margins(path: Path) -> None:
    """Зеркальные поля + нестандартные поля (подв. камень №18)."""
    doc = Document()
    s = doc.sections[0]
    s.left_margin = Mm(10)
    s.right_margin = Mm(10)
    s.top_margin = Mm(10)
    s.bottom_margin = Mm(10)
    doc.settings.element.append(OxmlElement("w:mirrorMargins"))
    doc.add_paragraph("Документ с зеркальными полями и переплётом.")
    doc.save(str(path))


def unmarked_structure(path: Path) -> None:
    """Неразмеченная структура как в реальных ВКР (для Фазы 2).

    Кодирует реальные паттерны и ЛОЖНЫЕ срабатывания, найденные на 5 настоящих ВКР:
    структурные элементы и нумерованные заголовки под стилем Normal, плюс
    TOC-строки с лидерами, поля-задания и пункты списков, которые трогать нельзя.
    """
    doc = Document()

    # TOC-строки (точечные лидеры + номер страницы) — НЕ заголовки, защитить.
    _set_run(doc.add_paragraph().add_run("ВВЕДЕНИЕ ............................. 3"), "Calibri", 12)
    _set_run(doc.add_paragraph().add_run("1.1 История вопроса ................. 5"), "Calibri", 12)

    # Структурный элемент как Normal (должен распознаться).
    _set_run(doc.add_paragraph().add_run("ВВЕДЕНИЕ"), "Calibri", 14)
    _set_run(doc.add_paragraph().add_run("Текст введения про актуальность темы исследования."), "Calibri", 12)

    # Нумерованные многоуровневые заголовки как Normal (должны распознаться).
    _set_run(doc.add_paragraph().add_run("1.1 История возникновения вопроса"), "Calibri", 12)
    _set_run(doc.add_paragraph().add_run("Содержательный абзац подраздела."), "Calibri", 12)
    _set_run(doc.add_paragraph().add_run("1.1.1 Подробности и детали"), "Calibri", 12)

    # ЛОЖНЫЕ срабатывания — НЕ заголовки:
    # поле-задание (одноуровневое, длинное, двоеточие)
    _set_run(
        doc.add_paragraph().add_run(
            "1. Тема: разработка мероприятий по оптимизации производственного процесса предприятия"
        ),
        "Calibri",
        12,
    )
    # пункт списка (одноуровневый, длинный, предложение)
    _set_run(
        doc.add_paragraph().add_run(
            "1. Кубики Кооса. Ребенку дают 16 небольших кубиков и просят собрать узор по образцу."
        ),
        "Calibri",
        12,
    )

    # Подписи (должны распознаться как Caption).
    _set_run(doc.add_paragraph().add_run("Рисунок 1 — Схема процесса"), "Calibri", 12)
    _set_run(doc.add_paragraph().add_run("Рис. 1.1. Роль управления"), "Calibri", 12)
    _set_run(doc.add_paragraph().add_run("Таблица 2 — Основные показатели"), "Calibri", 12)
    # ЛОЖНОЕ срабатывание — ссылка на рисунок в тексте (после номера нет разделителя).
    _set_run(
        doc.add_paragraph().add_run(
            "Рисунок 1 показывает устойчивый рост ключевых показателей за пятилетний период."
        ),
        "Calibri",
        12,
    )

    # Ещё структурные элементы как Normal (в т.ч. в нижнем регистре — проверка caps).
    _set_run(doc.add_paragraph().add_run("Заключение"), "Calibri", 14)
    _set_run(doc.add_paragraph().add_run("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"), "Calibri", 14)
    doc.save(str(path))


def legacy_doc(path: Path) -> None:
    """Псевдо-.doc со старой OLE2-сигнатурой (для проверки ошибки формата)."""
    ole2_magic = bytes([0xD0, 0xCF, 0x11, 0xE0, 0xA1, 0xB1, 0x1A, 0xE1])
    path.write_bytes(ole2_magic + b"\x00" * 512)


def broken_docx(path: Path) -> None:
    """Файл с расширением .docx, не являющийся валидным zip/docx."""
    path.write_bytes(b"This is not a real .docx file, just garbage bytes.")


# (имя файла, builder). already_gost и ошибки обрабатываются отдельно.
BUILDERS = {
    "calibri_default.docx": calibri_default,
    "direct_formatting.docx": direct_formatting,
    "multi_section.docx": multi_section,
    "styled_headings.docx": styled_headings,
    "with_table.docx": with_table,
    "with_image.docx": with_image,
    "with_list.docx": with_list,
    "localized_styles.docx": localized_styles,
    "title_and_body.docx": title_and_body,
    "with_hyperlink.docx": with_hyperlink,
    "with_revisions.docx": with_revisions,
    "mirror_margins.docx": mirror_margins,
    "unmarked_structure.docx": unmarked_structure,
}


def generate_all(force: bool = False) -> None:
    """Сгенерировать все фикстуры в FIXTURES_DIR."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    for filename, builder in BUILDERS.items():
        target = FIXTURES_DIR / filename
        if force or not target.exists():
            builder(target)
    # Ошибочные файлы
    legacy_doc(FIXTURES_DIR / "legacy.doc")
    broken_docx(FIXTURES_DIR / "broken.docx")
    # already_gost — только если форматтер готов
    already_gost(FIXTURES_DIR / "already_gost.docx")


if __name__ == "__main__":
    generate_all(force=True)
    print(f"Фикстуры записаны в {FIXTURES_DIR}")
