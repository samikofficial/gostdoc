"""Поля всех секций и сквозная нумерация страниц (ТЗ, разд. 7.5).

Поля и позиция номера страницы берутся из профиля (дефолт — ГОСТ 7.32-2017).

Риски (ТЗ, разд. 8):
- №3 несколько секций → обходим document.sections целиком, а не [0].
- №4 поля PAGE нет в API → вставляем через xml_utils.
- №5 титул без номера, но в сквозной нумерации → different first page + пустой
  первый колонтитул, поле PAGE в основном; старт нумерации не сбрасываем.
- №14 идемпотентность → add_page_number_field не дублирует поле.
- №18 зеркальные поля/переплёт → снимаем mirrorMargins и обнуляем gutter.
"""

from __future__ import annotations

from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from . import xml_utils
from .profile import GOST, PageNumberPosition, Profile

_TOP_POSITIONS = (PageNumberPosition.TOP_RIGHT,)


def _disable_mirror_margins(document: _Document) -> None:
    settings_el = document.settings.element
    mirror = settings_el.find(qn("w:mirrorMargins"))
    if mirror is not None:
        settings_el.remove(mirror)


def normalize_margins(document: _Document, profile: Profile = GOST) -> None:
    """Выставить поля профиля каждой секции; снять зеркальные поля и переплёт."""
    _disable_mirror_margins(document)
    for section in document.sections:
        section.left_margin = profile.margin_left
        section.right_margin = profile.margin_right
        section.top_margin = profile.margin_top
        section.bottom_margin = profile.margin_bottom
        section.gutter = 0


def _set_page_field(container, alignment) -> None:
    para = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    para.alignment = alignment
    xml_utils.add_page_number_field(para)


def setup_page_numbering(document: _Document, profile: Profile = GOST) -> None:
    """Сквозная нумерация по позиции из профиля; на первой странице номер не печатается."""
    if profile.page_number == PageNumberPosition.NONE:
        return
    sections = document.sections
    if not sections:
        return

    top = profile.page_number in _TOP_POSITIONS
    alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if profile.page_number == PageNumberPosition.BOTTOM_CENTER
        else WD_ALIGN_PARAGRAPH.RIGHT
    )

    first = sections[0]
    first.different_first_page_header_footer = True
    container = first.header if top else first.footer
    container.is_linked_to_previous = False  # создаём собственный колонтитул
    _set_page_field(container, alignment)
    # Остальные секции наследуют колонтитул — нумерация остаётся сквозной.
    for section in sections[1:]:
        (section.header if top else section.footer).is_linked_to_previous = True
