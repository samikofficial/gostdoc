"""Конвейер: открыть → привести оформление к ГОСТ → сохранить (ТЗ, разд. 7.5).

Главный инвариант: меняем только оформление. Текст тела (абзацы + ячейки таблиц)
до и после обработки совпадает побайтно; единственный добавленный текст — поле PAGE
в колонтитуле (в тело не входит).

Риски (ТЗ, разд. 8):
- №12,13 не-.docx/.doc/повреждён/запаролен → понятная ошибка, без трассировки.
- полнота обхода → рекурсивно обходим тело, включая вложенные таблицы.
- №16 рисунки в run'ах → run'ы не пересоздаём (см. runs.py).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table

from . import classify, constants as c, detect, paragraphs, runs, sections, styles
from .profile import GOST, PageNumberPosition, Profile
from .xml_utils import lengths_equal, paragraph_has_page_field

_EMU_PER_MM = 36000

_OLE2_MAGIC = b"\xd0\xcf\x11\xe0"
_ZIP_MAGIC = b"PK"

# Теги исправлений (режим рецензирования) — их не принимаем/не отклоняем, но предупреждаем.
_REVISION_TAGS = ("w:ins", "w:del", "w:moveFrom", "w:moveTo")


class GostDocError(Exception):
    """Понятная пользователю ошибка (плохой формат, повреждённый/запароленный файл)."""


def _validate_and_open(input_path: str) -> _Document:
    path = Path(input_path)
    if not path.exists():
        raise GostDocError(f"Файл не найден: {input_path}")

    ext = path.suffix.lower()
    if ext == ".doc":
        raise GostDocError("Формат .doc не поддерживается, пересохраните файл в .docx.")
    if ext != ".docx":
        raise GostDocError("Это не файл .docx.")

    signature = path.read_bytes()[:4]
    if signature == _OLE2_MAGIC:
        raise GostDocError(
            "Файл .docx в двоичном формате (повреждён или защищён паролем). "
            "Снимите защиту и пересохраните."
        )
    if signature[:2] != _ZIP_MAGIC:
        raise GostDocError("Файл повреждён или не является документом .docx.")

    try:
        return Document(str(path))
    except Exception as exc:  # noqa: BLE001 — переводим в понятную ошибку, не глотаем
        raise GostDocError(
            f"Не удалось открыть документ (повреждён или защищён паролем): {input_path}"
        ) from exc


def _normalize_runs_in_paragraph(paragraph, category: str, profile: Profile) -> None:
    is_heading = category == classify.CATEGORY_HEADING
    # Включаем run'ы внутри гиперссылок (подв. камень №17): иначе они остаются Calibri.
    run_list = list(paragraph.runs)
    for hyperlink in paragraph.hyperlinks:
        run_list.extend(hyperlink.runs)
    for run in run_list:
        if is_heading:
            runs.normalize_heading_run(run, bold=profile.bold_headings)
        else:
            runs.normalize_run(run)


def _process_paragraph(paragraph, in_table: bool, profile: Profile) -> None:
    category = classify.classify_paragraph(paragraph, in_table)
    paragraphs.apply_paragraph_format(paragraph, category)
    _normalize_runs_in_paragraph(paragraph, category, profile)


def _process_table(table: Table, profile: Profile) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _process_paragraph(paragraph, in_table=True, profile=profile)
            for nested in cell.tables:  # вложенные таблицы
                _process_table(nested, profile)


def _normalize_body(document: _Document, profile: Profile) -> None:
    for paragraph in document.paragraphs:
        _process_paragraph(paragraph, in_table=False, profile=profile)
    for table in document.tables:
        _process_table(table, profile)


def _normalize_header_footer_fonts(document: _Document, profile: Profile) -> None:
    """Гарнитуру/кегль колонтитулов (включая поле PAGE) — к ГОСТ; формат не трогаем."""
    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    runs.normalize_run(run, font_size=profile.page_number_size)


def extract_body_text(document: _Document) -> str:
    """Чистый текст тела: абзацы документа + абзацы ячеек таблиц (без колонтитулов).

    Общий хелпер для теста инварианта и проверок (подводный камень №15).
    """
    parts: list[str] = []

    def collect_table(table: Table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    parts.append(paragraph.text)
                for nested in cell.tables:
                    collect_table(nested)

    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        collect_table(table)
    return "\n".join(parts)


def detect_warnings(document: _Document) -> list[str]:
    """Неблокирующие предупреждения: исправления и комментарии (подв. камень №19)."""
    warnings: list[str] = []
    root = document.element
    if any(root.findall(".//" + qn(tag)) for tag in _REVISION_TAGS):
        warnings.append(
            "Предупреждение: в документе есть исправления (режим рецензирования). Оформление "
            "приведено к ГОСТ, но итог может выглядеть иначе после принятия/отклонения правок."
        )
    if root.findall(".//" + qn("w:commentReference")):
        warnings.append("Предупреждение: в документе есть комментарии — они сохранены без изменений.")
    return warnings


def format_document(
    input_path: str,
    output_path: str,
    detect_structure: bool = False,
    profile: Profile = GOST,
) -> list[str]:
    """Привести оформление input_path к ГОСТ 7.32-2017 и сохранить в output_path.

    detect_structure=True включает Фазу 2 (распознавание неразмеченных заголовков/
    структурных элементов) — обратимую и логируемую.
    profile задаёт переопределения под методичку (поля, номер страницы, полужирность).

    Возвращает список сообщений: предупреждения (исправления/комментарии) и, если
    включена Фаза 2, лог авторазметки.
    """
    document = _validate_and_open(input_path)
    messages = detect_warnings(document)
    if detect_structure:
        marked = detect.detect_and_mark(document)
        if marked:
            messages.append(f"Фаза 2: размечено элементов структуры — {len(marked)}.")
            messages.extend(f"  {line}" for line in marked)
    styles.normalize_styles(document, profile)
    sections.normalize_margins(document, profile)
    sections.setup_page_numbering(document, profile)
    _normalize_body(document, profile)
    _normalize_header_footer_fonts(document, profile)
    document.save(output_path)
    return messages


def _mm(value: int) -> int:
    return round(int(value) / _EMU_PER_MM)


def check_compliance(document: _Document, profile: Profile = GOST) -> list[str]:
    """Вернуть список несоответствий профилю (пусто = соответствует).

    Проверяем верифицируемые ключевые параметры (поля, шрифт Normal, нумерацию).
    """
    problems: list[str] = []
    expected = (
        ("левое", "left_margin", profile.margin_left),
        ("правое", "right_margin", profile.margin_right),
        ("верхнее", "top_margin", profile.margin_top),
        ("нижнее", "bottom_margin", profile.margin_bottom),
    )
    for i, section in enumerate(document.sections, start=1):
        for label, attr, want in expected:
            if not lengths_equal(getattr(section, attr), want):
                problems.append(f"Секция {i}: {label} поле не равно {_mm(want)} мм.")

    normal = document.styles["Normal"]
    if normal.font.name != c.FONT_NAME:
        problems.append("Стиль «Обычный»: шрифт не Times New Roman.")
    if normal.font.size != c.FONT_SIZE_BODY:
        problems.append("Стиль «Обычный»: кегль не равен 14 пт.")

    if document.sections and profile.page_number != PageNumberPosition.NONE:
        first = document.sections[0]
        top = profile.page_number == PageNumberPosition.TOP_RIGHT
        container = first.header if top else first.footer
        if not first.different_first_page_header_footer:
            problems.append("Не включён особый колонтитул первой страницы (номер на титуле).")
        if not any(paragraph_has_page_field(p) for p in container.paragraphs):
            problems.append("В колонтитуле нет поля номера страницы.")
    return problems


def check_file(input_path: str, profile: Profile = GOST) -> list[str]:
    """Открыть документ (с валидацией формата) и проверить соответствие профилю."""
    return check_compliance(_validate_and_open(input_path), profile)
